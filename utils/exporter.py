"""
Report Export Utilities
Handles Word document generation and PDF conversion
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from typing import Dict, List, Optional
import os
from pathlib import Path


class ReportGenerator:
    """Generates professional BVA reports in Word and PDF formats"""
    
    def __init__(self, output_dir: str = "reports"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
    def create_bva_report(self, 
                         client_data: Dict,
                         research_data: Dict,
                         calculations: Dict,
                         recommendations: List[str]) -> str:
        """
        Create a comprehensive BVA report
        
        Returns:
            Path to the generated Word document
        """
        doc = Document()
        
        # Set up document styles
        self._setup_styles(doc)
        
        # Add content sections
        self._add_cover_page(doc, client_data)
        self._add_executive_summary(doc, client_data, calculations)
        self._add_company_overview(doc, research_data)
        self._add_business_priorities(doc, research_data)
        self._add_value_analysis(doc, calculations)
        self._add_detailed_calculations(doc, calculations)
        self._add_recommendations(doc, recommendations)
        self._add_appendix(doc, research_data, calculations)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"BVA_{client_data['name'].replace(' ', '_')}_{timestamp}.docx"
        filepath = self.output_dir / filename
        
        doc.save(str(filepath))
        print(f"✅ Report generated: {filepath}")
        
        return str(filepath)
    
    def _setup_styles(self, doc: Document):
        """Configure document styles"""
        styles = doc.styles
        
        # Title style
        if 'BVA Title' not in styles:
            title_style = styles.add_style('BVA Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Arial'
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(15, 98, 254)  # IBM Blue
        
        # Heading styles
        if 'BVA Heading 1' not in styles:
            h1_style = styles.add_style('BVA Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.name = 'Arial'
            h1_style.font.size = Pt(20)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(15, 98, 254)
        
        if 'BVA Heading 2' not in styles:
            h2_style = styles.add_style('BVA Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.name = 'Arial'
            h2_style.font.size = Pt(16)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(22, 22, 22)
    
    def _add_cover_page(self, doc: Document, client_data: Dict):
        """Add cover page"""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("Business Value Assessment")
        title_run.font.size = Pt(32)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(15, 98, 254)
        
        doc.add_paragraph()  # Spacing
        
        # Client name
        client = doc.add_paragraph()
        client.alignment = WD_ALIGN_PARAGRAPH.CENTER
        client_run = client.add_run(client_data.get('name', 'Client Name'))
        client_run.font.size = Pt(24)
        client_run.font.bold = True
        
        doc.add_paragraph()  # Spacing
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(f"{client_data.get('product', 'IBM Solution')} Implementation")
        subtitle_run.font.size = Pt(18)
        
        doc.add_paragraph()  # Spacing
        
        # Date
        date = doc.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date.add_run(datetime.now().strftime("%B %d, %Y"))
        date_run.font.size = Pt(14)
        
        # Page break
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, client_data: Dict, calculations: Dict):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        summary = calculations.get('summary', {})
        roi = calculations.get('roi', {})
        payback = calculations.get('payback', {})
        
        # Key findings
        doc.add_paragraph(
            f"This Business Value Assessment analyzes the potential impact of implementing "
            f"{client_data.get('product', 'the proposed solution')} at {client_data.get('name', 'the organization')}. "
            f"Our analysis reveals significant opportunities for value creation across efficiency, "
            f"risk mitigation, and growth enablement."
        )
        
        doc.add_paragraph()
        
        # Key metrics table
        doc.add_heading('Key Financial Metrics', level=2)
        
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        metrics = [
            ('Total Annual Value', f"${summary.get('total_annual_value', 0):,.0f}"),
            ('3-Year Total Value', f"${summary.get('total_value_over_period', 0):,.0f}"),
            ('Project Investment', f"${summary.get('project_cost', 0):,.0f}"),
            ('Return on Investment', f"{roi.get('roi_percentage', 0):.1f}%"),
            ('Payback Period', f"{payback.get('payback_months', 0):.1f} months")
        ]
        
        for i, (label, value) in enumerate(metrics):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Bold the labels
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_page_break()
    
    def _add_company_overview(self, doc: Document, research_data: Dict):
        """Add company overview section"""
        doc.add_heading('Company Overview', level=1)
        
        company_info = research_data.get('company_info', {})
        
        doc.add_paragraph(
            f"{research_data.get('company_name', 'The company')} operates in the "
            f"{research_data.get('industry', 'industry')} sector. "
            f"{company_info.get('description', 'The organization is focused on delivering value to its customers.')}"
        )
        
        doc.add_paragraph()
        
        # Industry context
        doc.add_heading('Industry Context', level=2)
        benchmarks = research_data.get('industry_benchmarks', {})
        
        if benchmarks:
            doc.add_paragraph(
                f"The {research_data.get('industry', 'industry')} industry is characterized by "
                f"rapid technological change and increasing competitive pressure. Key industry trends include:"
            )
            
            pain_points = benchmarks.get('common_pain_points', [])
            for point in pain_points:
                doc.add_paragraph(f"• {point}", style='List Bullet')
    
    def _add_business_priorities(self, doc: Document, research_data: Dict):
        """Add business priorities section"""
        doc.add_heading('Strategic Business Priorities', level=1)
        
        doc.add_paragraph(
            "Based on our research of public filings, investor relations materials, and market analysis, "
            "we have identified the following strategic priorities:"
        )
        
        doc.add_paragraph()
        
        value_pillars = research_data.get('value_pillars', [])
        
        for i, pillar in enumerate(value_pillars[:5], 1):
            doc.add_heading(f"{i}. {pillar.get('priority', 'Priority')}", level=2)
            doc.add_paragraph(
                f"Source: {pillar.get('source', 'Research')} | "
                f"Relevance Score: {pillar.get('relevance_score', 0):.0%}"
            )
            doc.add_paragraph()
        
        doc.add_page_break()
    
    def _add_value_analysis(self, doc: Document, calculations: Dict):
        """Add value analysis section"""
        doc.add_heading('Value Analysis', level=1)
        
        value_breakdown = calculations.get('value_breakdown', {})
        
        doc.add_paragraph(
            "Our analysis identifies value creation across three key dimensions:"
        )
        
        doc.add_paragraph()
        
        # Efficiency savings
        doc.add_heading('Efficiency Savings', level=2)
        efficiency = value_breakdown.get('efficiency_savings', 0)
        doc.add_paragraph(
            f"Annual Value: ${efficiency:,.0f}\n\n"
            f"Through automation and process optimization, the solution will reduce manual effort, "
            f"eliminate redundant tasks, and accelerate workflows."
        )
        
        doc.add_paragraph()
        
        # Risk mitigation
        doc.add_heading('Risk Mitigation', level=2)
        risk = value_breakdown.get('risk_mitigation', 0)
        doc.add_paragraph(
            f"Annual Value: ${risk:,.0f}\n\n"
            f"By reducing downtime, improving compliance, and enhancing security, the solution "
            f"mitigates significant operational and financial risks."
        )
        
        doc.add_paragraph()
        
        # Growth enablement
        doc.add_heading('Growth Enablement', level=2)
        growth = value_breakdown.get('growth_enablement', 0)
        doc.add_paragraph(
            f"Annual Value: ${growth:,.0f}\n\n"
            f"The solution enables faster time-to-market, improved customer experience, "
            f"and new revenue opportunities."
        )
        
        doc.add_page_break()
    
    def _add_detailed_calculations(self, doc: Document, calculations: Dict):
        """Add detailed calculations section"""
        doc.add_heading('Detailed Calculations', level=1)
        
        calc_details = calculations.get('calculation_details', [])
        
        for i, calc in enumerate(calc_details, 1):
            if 'error' in calc:
                continue
                
            doc.add_heading(f"{i}. {calc.get('formula_name', 'Calculation')}", level=2)
            
            # Variables used
            doc.add_paragraph("Input Variables:", style='Heading 3')
            variables = calc.get('variables_used', {})
            for var_name, var_value in variables.items():
                doc.add_paragraph(f"• {var_name}: {var_value}", style='List Bullet')
            
            doc.add_paragraph()
            
            # Formula
            doc.add_paragraph(f"Formula: {calc.get('formula', 'N/A')}")
            
            doc.add_paragraph()
            
            # Result
            result_para = doc.add_paragraph()
            result_run = result_para.add_run(
                f"Result: ${calc.get('result', 0):,.2f} {calc.get('unit', '')}"
            )
            result_run.font.bold = True
            result_run.font.size = Pt(14)
            
            doc.add_paragraph()
    
    def _add_recommendations(self, doc: Document, recommendations: List[str]):
        """Add recommendations section"""
        doc.add_heading('Recommendations', level=1)
        
        doc.add_paragraph(
            "Based on our analysis, we recommend the following actions:"
        )
        
        doc.add_paragraph()
        
        for i, rec in enumerate(recommendations, 1):
            doc.add_heading(f"{i}. {rec}", level=2)
            doc.add_paragraph()
        
        doc.add_page_break()
    
    def _add_appendix(self, doc: Document, research_data: Dict, calculations: Dict):
        """Add appendix section"""
        doc.add_heading('Appendix', level=1)
        
        # Methodology
        doc.add_heading('A. Methodology', level=2)
        doc.add_paragraph(
            "This Business Value Assessment was conducted using a comprehensive, multi-source "
            "research approach combined with industry-standard financial modeling techniques."
        )
        
        doc.add_paragraph()
        
        # Data sources
        doc.add_heading('B. Data Sources', level=2)
        doc.add_paragraph("• SEC Filings (10-K, 10-Q)", style='List Bullet')
        doc.add_paragraph("• Investor Relations Materials", style='List Bullet')
        doc.add_paragraph("• Industry Benchmarks (Gartner, IDC)", style='List Bullet')
        doc.add_paragraph("• Market Analysis and News", style='List Bullet')
        
        doc.add_paragraph()
        
        # Assumptions
        doc.add_heading('C. Key Assumptions', level=2)
        summary = calculations.get('summary', {})
        doc.add_paragraph(
            f"• Analysis Period: {summary.get('time_period_years', 3)} years"
        )
        doc.add_paragraph("• Discount Rate: 10%")
        doc.add_paragraph("• Risk Adjustment Factor: 85%")
    
    def convert_to_pdf(self, docx_path: str) -> str:
        """
        Convert Word document to PDF
        
        Note: This requires additional dependencies or external tools
        For production, consider using:
        - docx2pdf library (Windows/Mac)
        - LibreOffice in headless mode (cross-platform)
        - Cloud conversion services
        """
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        try:
            # Attempt conversion using docx2pdf (if available)
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                print(f"✅ PDF generated: {pdf_path}")
                return pdf_path
            except ImportError:
                print("⚠️  docx2pdf not available. Install with: pip install docx2pdf")
                print("   Alternatively, manually convert the Word document to PDF")
                return docx_path
                
        except Exception as e:
            print(f"⚠️  PDF conversion failed: {e}")
            print(f"   Word document available at: {docx_path}")
            return docx_path


def test_exporter():
    """Test the report generator"""
    generator = ReportGenerator()
    
    # Sample data
    client_data = {
        'name': 'Acme Corporation',
        'industry': 'Manufacturing',
        'product': 'IBM watsonx'
    }
    
    research_data = {
        'company_name': 'Acme Corporation',
        'industry': 'Manufacturing',
        'company_info': {
            'description': 'Leading manufacturer of industrial equipment'
        },
        'value_pillars': [
            {'priority': 'Digital Transformation', 'source': 'SEC Filings', 'relevance_score': 0.9},
            {'priority': 'Operational Efficiency', 'source': 'CEO Letter', 'relevance_score': 0.85}
        ],
        'industry_benchmarks': {
            'common_pain_points': [
                'Legacy system modernization',
                'Supply chain optimization',
                'Quality control automation'
            ]
        }
    }
    
    calculations = {
        'summary': {
            'total_annual_value': 1500000,
            'total_value_over_period': 4500000,
            'project_cost': 500000,
            'time_period_years': 3
        },
        'value_breakdown': {
            'efficiency_savings': 800000,
            'risk_mitigation': 400000,
            'growth_enablement': 300000
        },
        'roi': {
            'roi_percentage': 800,
            'net_value': 4000000
        },
        'payback': {
            'payback_months': 4.0
        },
        'calculation_details': [
            {
                'formula_name': 'Time Savings',
                'formula': '(manual_hours * hourly_rate) - (automated_hours * hourly_rate)',
                'variables_used': {'manual_hours': 1000, 'automated_hours': 200, 'hourly_rate': 50},
                'result': 40000,
                'unit': 'USD annually'
            }
        ]
    }
    
    recommendations = [
        'Proceed with Implementation',
        'Establish Success Metrics',
        'Plan Change Management'
    ]
    
    # Generate report
    report_path = generator.create_bva_report(
        client_data,
        research_data,
        calculations,
        recommendations
    )
    
    print(f"\n✅ Test report generated successfully!")
    print(f"📄 Location: {report_path}")


if __name__ == "__main__":
    test_exporter()

# Made with Bob
